#!/usr/bin/env python3
"""Generate the Beauty Manager Pro user manual as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2C, 0x24, 0x20)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0xB8, 0x95, 0x6A)
    h.font.name = 'Calibri'

style_toc = doc.styles.add_style('TOCEntry', WD_STYLE_TYPE.PARAGRAPH)
style_toc.font.name = 'Calibri'
style_toc.font.size = Pt(11)
style_toc.paragraph_format.space_after = Pt(2)

style_tip = doc.styles.add_style('Tip', WD_STYLE_TYPE.PARAGRAPH)
style_tip.font.name = 'Calibri'
style_tip.font.size = Pt(10)
style_tip.font.italic = True
style_tip.font.color.rgb = RGBColor(0x55, 0x80, 0x55)

style_note = doc.styles.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
style_note.font.name = 'Calibri'
style_note.font.size = Pt(10)
style_note.font.italic = True
style_note.font.color.rgb = RGBColor(0x99, 0x66, 0x33)


def add_tip(text):
    p = doc.add_paragraph(f"💡 {text}", style='Tip')
    return p

def add_note(text):
    p = doc.add_paragraph(f"⚠️ {text}", style='Note')
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')

def add_table_2col(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = headers[0]
    hdr[1].text = headers[1]
    for r in table.rows[0].cells:
        for p in r.paragraphs:
            p.runs[0].font.bold = True
    for i, (c1, c2) in enumerate(rows):
        table.rows[i + 1].cells[0].text = c1
        table.rows[i + 1].cells[1].text = c2
    doc.add_paragraph()

def add_table_3col(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        for p in table.rows[0].cells[j].paragraphs:
            p.runs[0].font.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph()

def section_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Beauty Manager Pro')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xB8, 0x95, 0x6A)
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Manuale Utente Completo')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2C, 0x24, 0x20)
run.font.name = 'Calibri'

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Versione 0.1.0 — 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

section_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Indice', level=1)
toc_items = [
    '1. Introduzione',
    '2. Primo Avvio e Configurazione',
    '3. Dashboard — Panoramica Giornaliera',
    '4. Agenda — Gestione Appuntamenti',
    '5. Clienti — Anagrafica e Segmentazione',
    '6. Operatrici — Gestione Staff',
    '7. Trattamenti — Catalogo Servizi',
    '8. Magazzino — Gestione Prodotti e Inventario',
    '9. Comunicazioni — Marketing e Template',
    '10. Report — Analisi Dati',
    '11. Insights — Business Intelligence',
    '12. Impostazioni — Configurazione',
    '13. Aggiornamenti — OTA Updates',
    '14. Scorciatoie e Trucchi',
]
for item in toc_items:
    doc.add_paragraph(item, style='TOCEntry')

section_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Introduzione', level=1)
doc.add_paragraph(
    'Beauty Manager Pro è un gestionale completo per centri estetici, sviluppato con tecnologia '
    'moderna (Tauri + React) per garantire prestazioni elevate e un\'esperienza utente fluida. '
    'Il software funziona interamente in locale sul tuo computer, senza necessità di connessione internet '
    'per le operazioni quotidiane.'
)
doc.add_paragraph()
doc.add_heading('Cosa puoi fare con Beauty Manager Pro', level=2)
add_bullet('Gestire l\'agenda appuntamenti con vista giorno, settimana e mese')
add_bullet('Registrare e segmentare automaticamente i clienti (VIP, Abituali, A Rischio, ecc.)')
add_bullet('Gestire il catalogo trattamenti con prezzi e durate')
add_bullet('Monitorare il magazzino prodotti con alert scorte e scadenze')
add_bullet('Inviare promemoria e auguri di compleanno via WhatsApp ed email')
add_bullet('Analizzare i dati con report dettagliati e insights intelligenti')
add_bullet('Personalizzare tema, colori e layout della dashboard')

doc.add_paragraph()
doc.add_heading('Requisiti di sistema', level=2)
add_table_2col(
    ['Requisito', 'Dettaglio'],
    [
        ('Sistema operativo', 'Windows 10/11 (64-bit) o macOS 11+'),
        ('RAM', 'Minimo 4 GB (consigliati 8 GB)'),
        ('Spazio disco', 'Circa 150 MB per l\'applicazione'),
        ('Schermo', 'Risoluzione minima 1280×720'),
    ]
)

section_break()

# ═══════════════════════════════════════════════════════════════
# 2. PRIMO AVVIO E CONFIGURAZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Primo Avvio e Configurazione', level=1)
doc.add_paragraph(
    'Al primo avvio, Beauty Manager Pro crea automaticamente il database locale e ti guida '
    'nella configurazione iniziale.'
)
doc.add_paragraph()

doc.add_heading('2.1 Installazione', level=2)
add_numbered('Scarica il file di installazione dalla pagina Releases di GitHub')
add_numbered('Windows: esegui il file .exe o .msi e segui la procedura guidata')
add_numbered('macOS: apri il file .dmg e trascina l\'app nella cartella Applicazioni')
add_numbered('Avvia Beauty Manager Pro dal menu Start o dal Launchpad')

doc.add_paragraph()
doc.add_heading('2.2 Configurazione iniziale consigliata', level=2)
doc.add_paragraph(
    'Prima di iniziare a lavorare, configura questi elementi fondamentali nell\'ordine indicato:'
)
add_numbered('Impostazioni → Dati Azienda: inserisci nome, indirizzo, orari di apertura/chiusura e giorni lavorativi')
add_numbered('Operatrici: crea almeno un\'operatrice (necessaria per l\'agenda)')
add_numbered('Trattamenti: crea le categorie e i trattamenti offerti con prezzi e durate')
add_numbered('Prodotti (opzionale): inserisci i prodotti usati nei trattamenti per il tracciamento magazzino')
add_numbered('Impostazioni → Email SMTP (opzionale): configura l\'invio email per promemoria e auguri')

add_note('Senza operatrici configurate, l\'agenda e il calcolo degli slot liberi non funzioneranno correttamente.')

section_break()

# ═══════════════════════════════════════════════════════════════
# 3. DASHBOARD
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Dashboard — Panoramica Giornaliera', level=1)
doc.add_paragraph(
    'La Dashboard è la prima schermata che vedi all\'apertura dell\'app. Offre una visione '
    'immediata dello stato della giornata e delle metriche chiave del tuo centro.'
)

doc.add_paragraph()
doc.add_heading('3.1 Card Principali', level=2)

doc.add_heading('Appuntamenti Oggi', level=3)
doc.add_paragraph(
    'Mostra il numero totale di appuntamenti della giornata, suddivisi per stato:'
)
add_table_2col(
    ['Stato', 'Significato'],
    [
        ('Confermati', 'Appuntamenti confermati dal cliente'),
        ('In attesa', 'Appuntamenti prenotati ma non ancora confermati'),
        ('In ritardo', 'Appuntamenti con orario già passato ma non ancora iniziati'),
        ('Completati', 'Appuntamenti conclusi'),
    ]
)
add_tip('Cliccando sulla card si apre direttamente l\'Agenda sulla data odierna.')

doc.add_heading('Fatturato Oggi', level=3)
doc.add_paragraph(
    'Mostra l\'incasso della giornata in euro, con due indicatori di tendenza:'
)
add_bullet('Freccia verde/rossa rispetto a ieri')
add_bullet('Freccia verde/rossa rispetto allo stesso giorno della settimana scorsa')

doc.add_heading('Prossimo Cliente', level=3)
doc.add_paragraph(
    'Mostra i dettagli del prossimo appuntamento in programma:'
)
add_bullet('Tempo rimanente (ore e minuti)')
add_bullet('Nome del cliente')
add_bullet('Trattamento prenotato')
add_bullet('Operatrice assegnata')
add_tip('Se non ci sono appuntamenti futuri, la card mostrerà "Nessun appuntamento in programma".')

doc.add_heading('Slot Liberi', level=3)
doc.add_paragraph(
    'Mostra quanti appuntamenti possono ancora essere inseriti nella giornata. '
    'Il calcolo si basa su:'
)
add_bullet('Tempo rimanente da adesso fino all\'orario di chiusura')
add_bullet('Numero di operatrici attive (ogni operatrice ha i propri slot)')
add_bullet('Durata media dei trattamenti nel catalogo')
add_bullet('Appuntamenti futuri già prenotati')
add_note('Se non ci sono operatrici configurate, il valore sarà sempre 0. Dopo l\'orario di chiusura, sarà 0.')

doc.add_paragraph()
doc.add_heading('3.2 Sezioni Aggiuntive', level=2)

doc.add_heading('Andamento Settimanale', level=3)
doc.add_paragraph(
    'Grafico che mostra il fatturato degli ultimi 7 giorni, utile per identificare i giorni più produttivi.'
)

doc.add_heading('Top 5 Trattamenti del Mese', level=3)
doc.add_paragraph(
    'Classifica dei 5 trattamenti che hanno generato più fatturato nel mese corrente, con barre proporzionali.'
)

doc.add_heading('Top 5 Clienti del Mese', level=3)
doc.add_paragraph(
    'Classifica dei 5 clienti che hanno speso di più nel mese corrente.'
)

doc.add_paragraph()
doc.add_heading('3.3 Personalizzazione Dashboard', level=2)
doc.add_paragraph(
    'Puoi personalizzare la dashboard cliccando sull\'icona ingranaggio in alto a destra:'
)
add_bullet('Nascondi/mostra sezioni con l\'icona occhio')
add_bullet('Riordina le sezioni trascinandole con il grip')
add_bullet('Ripristina il layout predefinito con il pulsante "Reset"')

section_break()

# ═══════════════════════════════════════════════════════════════
# 4. AGENDA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Agenda — Gestione Appuntamenti', level=1)
doc.add_paragraph(
    'L\'Agenda è il cuore operativo del gestionale. Permette di visualizzare, creare e '
    'gestire tutti gli appuntamenti con diverse modalità di vista.'
)

doc.add_paragraph()
doc.add_heading('4.1 Modalità di Visualizzazione', level=2)

doc.add_heading('Vista Giorno', level=3)
doc.add_paragraph(
    'La vista predefinita mostra la griglia oraria della giornata, con una colonna per ogni operatrice attiva. '
    'Ogni appuntamento è rappresentato da un blocco colorato che occupa lo spazio proporzionale alla sua durata.'
)
add_bullet('Le colonne rappresentano le operatrici')
add_bullet('Le righe rappresentano gli slot orari (configurabili nelle impostazioni)')
add_bullet('I blocchi colorati sono gli appuntamenti')

doc.add_heading('Vista Settimana', level=3)
doc.add_paragraph(
    'Mostra 7 giorni in un\'unica vista. Utile per avere una panoramica settimanale del carico di lavoro.'
)

doc.add_heading('Vista Mese', level=3)
doc.add_paragraph(
    'Calendario mensile tradizionale con indicatori sui giorni che hanno appuntamenti.'
)

doc.add_paragraph()
doc.add_heading('4.2 Navigazione', level=2)
add_bullet('Pulsanti freccia ← → per spostarsi al giorno/settimana/mese precedente o successivo')
add_bullet('Pulsante "Oggi" per tornare rapidamente alla data corrente')
add_bullet('La data mostrata cambia in base alla vista selezionata')

doc.add_paragraph()
doc.add_heading('4.3 Creare un Appuntamento', level=2)
add_numbered('Clicca su uno slot vuoto nella griglia (vista giorno) oppure sul pulsante "+ Nuovo"')
add_numbered('Si apre il modulo appuntamento con data e ora precompilati')
add_numbered('Seleziona il Cliente dal menu a tendina (puoi cercare per nome)')
add_numbered('Seleziona l\'Operatrice')
add_numbered('Seleziona il Trattamento — durata e prezzo si compilano automaticamente')
add_numbered('Modifica eventuali note o il prezzo applicato')
add_numbered('Clicca "Salva" per confermare')

add_tip('Se il cliente non esiste, puoi crearlo al volo: clicca "+ Nuovo cliente" nel menu a tendina '
        'e compila i dati essenziali. Il cliente verrà creato e selezionato automaticamente.')

doc.add_paragraph()
doc.add_heading('4.4 Modificare un Appuntamento', level=2)
doc.add_paragraph('Ci sono diversi modi per modificare un appuntamento:')
add_bullet('Clicca sull\'appuntamento per aprire il modulo di modifica')
add_bullet('Trascina l\'appuntamento per spostarlo a un altro orario (vista giorno)')
add_bullet('Trascina tra colonne per riassegnare a un\'altra operatrice')
add_bullet('Ridimensiona il blocco dal bordo inferiore per cambiare la durata')

doc.add_paragraph()
doc.add_heading('4.5 Stati dell\'Appuntamento', level=2)
add_table_2col(
    ['Stato', 'Descrizione'],
    [
        ('Prenotato', 'Appuntamento inserito, in attesa di conferma o della data'),
        ('In corso', 'Il trattamento è in corso'),
        ('Completato', 'Trattamento concluso — si attiva la sezione prodotti usati'),
        ('Annullato', 'Appuntamento cancellato — non occupa slot'),
        ('No Show', 'Il cliente non si è presentato — non occupa slot'),
    ]
)

doc.add_paragraph()
doc.add_heading('4.6 Prodotti Usati nel Trattamento', level=2)
doc.add_paragraph(
    'Quando imposti lo stato su "Completato", appare la sezione Prodotti Usati. '
    'Qui puoi registrare i prodotti consumati durante il trattamento:'
)
add_numbered('Clicca il pulsante "+" per aprire la ricerca prodotti')
add_numbered('Cerca per nome o scansiona un codice a barre')
add_numbered('Clicca sul prodotto per aggiungerlo con quantità 1')
add_numbered('Usa i pulsanti +/- per regolare la quantità')
add_numbered('Il magazzino viene aggiornato automaticamente al salvataggio (scarico uso)')

add_note('Se riporti lo stato da "Completato" a "Prenotato", i prodotti precedentemente scaricati vengono '
         'automaticamente reintegrati nel magazzino (movimento di reso).')

doc.add_paragraph()
doc.add_heading('4.7 Promemoria Appuntamento', level=2)
doc.add_paragraph(
    'Nel modulo appuntamento trovi il pulsante "Reminder" che permette di inviare un '
    'promemoria al cliente:'
)
add_bullet('WhatsApp: si apre WhatsApp con il messaggio precompilato (se il cliente ha il cellulare e il consenso)')
add_bullet('Email: si invia un\'email con il template configurato (se l\'email è presente e il consenso attivo)')
add_note('I canali disponibili dipendono dai dati di contatto e dai consensi del cliente.')

doc.add_paragraph()
doc.add_heading('4.8 Appuntamenti Passati Non Gestiti', level=2)
doc.add_paragraph(
    'Quando consulti un giorno passato che ha appuntamenti ancora in stato "Prenotato" o "In corso", '
    'appare automaticamente un popup che ti chiede di aggiornare lo stato:'
)
add_bullet('Completato: il trattamento è stato eseguito')
add_bullet('Annullato: l\'appuntamento è stato cancellato')
add_bullet('No Show: il cliente non si è presentato')
add_tip('Questo evita che rimangano appuntamenti "fantasma" nel sistema, falsando le statistiche.')

doc.add_paragraph()
doc.add_heading('4.9 Filtro Operatrici', level=2)
doc.add_paragraph(
    'Nella vista giorno puoi filtrare le operatrici visibili:'
)
add_bullet('Clicca sul filtro operatrici in alto')
add_bullet('Seleziona/deseleziona le singole operatrici')
add_bullet('"Seleziona tutte" per visualizzare tutto lo staff')

section_break()

# ═══════════════════════════════════════════════════════════════
# 5. CLIENTI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Clienti — Anagrafica e Segmentazione', level=1)
doc.add_paragraph(
    'La sezione Clienti gestisce l\'intero database clienti con ricerca, segmentazione '
    'automatica e profilo dettagliato per ciascun cliente.'
)

doc.add_paragraph()
doc.add_heading('5.1 Lista Clienti', level=2)
doc.add_paragraph(
    'La schermata è divisa in due parti: la lista a sinistra e il dettaglio a destra.'
)
add_bullet('Barra di ricerca in tempo reale: cerca per nome, cognome, cellulare o email')
add_bullet('I risultati si aggiornano mentre digiti')
add_bullet('Clicca su un cliente per visualizzarne il dettaglio')

doc.add_paragraph()
doc.add_heading('5.2 Filtri di Segmentazione', level=2)
doc.add_paragraph('Sopra la lista, i filtri rapidi permettono di segmentare i clienti:')
add_table_2col(
    ['Filtro', 'Descrizione'],
    [
        ('Tutti', 'Tutti i clienti attivi'),
        ('Nuovi', 'Clienti aggiunti negli ultimi 30 giorni'),
        ('Compleanni', 'Clienti con compleanno nei prossimi 14 giorni'),
        ('Inattivi', 'Clienti disattivati (visibili solo con questo filtro)'),
    ]
)

doc.add_paragraph()
doc.add_heading('5.3 Creare un Nuovo Cliente', level=2)
add_numbered('Clicca il pulsante "+ Nuovo" in alto a destra')
add_numbered('Compila i campi: Nome, Cognome (obbligatori), Cellulare, Email, Data di nascita')
add_numbered('Imposta i consensi: WhatsApp, Email, SMS, Marketing')
add_numbered('Aggiungi eventuali note')
add_numbered('Clicca "Salva"')

doc.add_paragraph()
doc.add_heading('5.4 Scheda Cliente — Panoramica', level=2)
doc.add_paragraph('Selezionando un cliente, il pannello destro mostra:')
add_bullet('Statistiche: numero visite, spesa totale, stato VIP')
add_bullet('Ultimi appuntamenti con data, trattamento e importo')
add_bullet('Badge "Nuovo" per clienti recenti')
add_bullet('Indicatore compleanno se imminente')

doc.add_paragraph()
doc.add_heading('5.5 Scheda Cliente — Anagrafica', level=2)
doc.add_paragraph(
    'Il tab Anagrafica permette di visualizzare e modificare tutti i dati personali:'
)
add_bullet('Dati personali: nome, cognome, data di nascita')
add_bullet('Contatti: cellulare, email')
add_bullet('Indirizzo: via, città, CAP, provincia')
add_bullet('Consensi: marketing, SMS, WhatsApp, email')
add_bullet('Note libere')
add_tip('Clicca "Modifica" per abilitare i campi, apporta le modifiche e clicca "Salva".')

doc.add_paragraph()
doc.add_heading('5.6 Scheda Cliente — Profilo Estetico', level=2)
doc.add_paragraph(
    'Il tab Estetica raccoglie informazioni utili per i trattamenti:'
)
add_bullet('Tipo di pelle')
add_bullet('Allergie e sensibilità')
add_bullet('Preferenze sui trattamenti')
add_bullet('Storico trattamenti effettuati')

doc.add_paragraph()
doc.add_heading('5.7 Disattivare / Eliminare un Cliente', level=2)
add_bullet('Disattivare: il cliente viene nascosto dalla lista ma i dati restano nel database. Può essere riattivato.')
add_bullet('Eliminare: rimozione definitiva e irreversibile di tutti i dati del cliente.')
add_note('L\'eliminazione è irreversibile. Preferisci la disattivazione se non sei sicuro.')

section_break()

# ═══════════════════════════════════════════════════════════════
# 6. OPERATRICI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Operatrici — Gestione Staff', level=1)
doc.add_paragraph(
    'La sezione Operatrici gestisce lo staff del centro estetico. Ogni operatrice ha '
    'una colonna dedicata nell\'agenda e un colore identificativo nel calendario.'
)

doc.add_paragraph()
doc.add_heading('6.1 Creare un\'Operatrice', level=2)
add_numbered('Clicca "+ Nuovo" in alto a destra')
add_numbered('Compila: Nome e Cognome (obbligatori)')
add_numbered('Il Codice (es. OP001) viene generato automaticamente')
add_numbered('Aggiungi: cellulare, email, specializzazioni, note')
add_numbered('Scegli il colore per il calendario dalla palette')
add_numbered('Clicca "Salva"')

doc.add_paragraph()
doc.add_heading('6.2 Dettaglio Operatrice', level=2)
doc.add_paragraph(
    'Selezionando un\'operatrice dal pannello sinistro, a destra vedi:'
)
add_bullet('Nome completo e codice')
add_bullet('Contatti (cellulare, email)')
add_bullet('Specializzazioni')
add_bullet('Colore calendario')
add_bullet('Stato attivo/inattivo')
add_bullet('Note')

doc.add_paragraph()
doc.add_heading('6.3 Disattivare un\'Operatrice', level=2)
doc.add_paragraph(
    'Un\'operatrice disattivata non appare nell\'agenda e non è selezionabile per nuovi appuntamenti, '
    'ma i suoi appuntamenti passati restano nel sistema.'
)

section_break()

# ═══════════════════════════════════════════════════════════════
# 7. TRATTAMENTI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Trattamenti — Catalogo Servizi', level=1)
doc.add_paragraph(
    'Il catalogo trattamenti definisce tutti i servizi offerti dal centro, organizzati per categorie.'
)

doc.add_paragraph()
doc.add_heading('7.1 Categorie', level=2)
doc.add_paragraph(
    'Prima di creare i trattamenti, definisci le categorie (es. Viso, Corpo, Epilazione, ecc.). '
    'Clicca "Gestisci categorie" per creare, rinominare o eliminare le categorie.'
)

doc.add_paragraph()
doc.add_heading('7.2 Creare un Trattamento', level=2)
add_numbered('Clicca "+ Nuovo"')
add_numbered('Compila: Nome (obbligatorio), Categoria (obbligatoria)')
add_numbered('Imposta la Durata in minuti (usata per calcolare gli slot in agenda)')
add_numbered('Imposta il Prezzo in euro (precompilato quando prenoti un appuntamento)')
add_numbered('Aggiungi eventuale descrizione e note operative')
add_numbered('Clicca "Salva"')

doc.add_paragraph()
doc.add_heading('7.3 Filtri', level=2)
add_bullet('Cerca per nome o descrizione')
add_bullet('Filtra per categoria con i chip colorati')
add_bullet('Visualizza tutti con il filtro "Tutti"')

doc.add_paragraph()
doc.add_heading('7.4 Dettaglio Trattamento', level=2)
doc.add_paragraph(
    'Selezionando un trattamento vedi: nome, categoria, prezzo, durata, descrizione e note operative. '
    'Puoi modificare o disattivare il trattamento dal menu azioni.'
)
add_tip('Un trattamento disattivato non appare nelle scelte per nuovi appuntamenti ma resta visibile '
        'nello storico.')

section_break()

# ═══════════════════════════════════════════════════════════════
# 8. MAGAZZINO
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Magazzino — Gestione Prodotti e Inventario', level=1)
doc.add_paragraph(
    'Il Magazzino gestisce l\'intero ciclo di vita dei prodotti: dall\'inserimento al consumo, '
    'con alert automatici per scorte basse e scadenze.'
)

doc.add_paragraph()
doc.add_heading('8.1 Tab Articoli — Lista Prodotti', level=2)
doc.add_paragraph(
    'La lista principale dei prodotti con filtri avanzati:'
)
add_bullet('Ricerca per nome, codice o codice a barre')
add_bullet('Filtro per categoria')
add_bullet('Filtro "Sotto scorta": mostra solo prodotti con giacenza inferiore al minimo')
add_bullet('Filtro "In scadenza": mostra prodotti che scadono entro 30 giorni')
add_bullet('Toggle "Mostra inattivi"')

doc.add_paragraph()
doc.add_heading('8.2 Creare un Prodotto', level=2)
add_numbered('Clicca "+ Nuovo"')
add_numbered('Compila i campi obbligatori: Nome, Categoria')
add_numbered('Campi opzionali ma consigliati:')
add_bullet('Codice prodotto e Codice a barre', level=1)
add_bullet('Marca e Linea', level=1)
add_bullet('Prezzo di acquisto e Prezzo di vendita', level=1)
add_bullet('Unità di misura e Capacità', level=1)
add_bullet('Giacenza iniziale', level=1)
add_bullet('Scorta minima (per gli alert) e Livello riordino', level=1)
add_bullet('Data di scadenza', level=1)
add_bullet('Tipo uso: interno, vendita, entrambi', level=1)
add_numbered('Clicca "Salva"')

doc.add_paragraph()
doc.add_heading('8.3 Scheda Prodotto', level=2)
doc.add_paragraph(
    'Selezionando un prodotto vedi tutti i dettagli: giacenza attuale, scorta minima, prezzi, '
    'scadenza e stato. Puoi modificare i dati in modalità modifica.'
)

doc.add_paragraph()
doc.add_heading('8.4 Tab Carico — Registrare Arrivo Merce', level=2)
doc.add_paragraph(
    'Ogni volta che ricevi prodotti dal fornitore:'
)
add_numbered('Vai al tab "Carico"')
add_numbered('Seleziona il prodotto')
add_numbered('Inserisci la quantità ricevuta')
add_numbered('Aggiungi eventuali note (es. numero bolla, fornitore)')
add_numbered('Conferma: la giacenza viene aggiornata automaticamente')

doc.add_paragraph()
doc.add_heading('8.5 Tab Scarico — Registrare Uscita Merce', level=2)
doc.add_paragraph(
    'Per registrare prodotti venduti o utilizzati manualmente (fuori dagli appuntamenti):'
)
add_numbered('Vai al tab "Scarico"')
add_numbered('Seleziona il prodotto')
add_numbered('Inserisci la quantità')
add_numbered('Seleziona il tipo: uso interno o vendita')
add_numbered('Conferma')
add_tip('I prodotti usati durante un appuntamento (sezione "Prodotti Usati") vengono scaricati '
        'automaticamente — non serve registrarli manualmente.')

doc.add_paragraph()
doc.add_heading('8.6 Tab Movimenti — Storico', level=2)
doc.add_paragraph(
    'Visualizza lo storico completo di tutti i movimenti di magazzino, filtrabile per periodo. '
    'Ogni movimento mostra: prodotto, tipo (carico/scarico/uso/reso), quantità, data e note.'
)

doc.add_paragraph()
doc.add_heading('8.7 Alert Automatici', level=2)
doc.add_paragraph(
    'In cima alla pagina, un banner giallo ti avvisa quando:'
)
add_bullet('Ci sono prodotti con giacenza sotto la scorta minima')
add_bullet('Ci sono prodotti in scadenza entro 30 giorni')
add_tip('Clicca sull\'alert per filtrare automaticamente i prodotti interessati.')

section_break()

# ═══════════════════════════════════════════════════════════════
# 9. COMUNICAZIONI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Comunicazioni — Marketing e Template', level=1)
doc.add_paragraph(
    'La sezione Comunicazioni permette di gestire i template per i messaggi WhatsApp ed email, '
    'e di inviare auguri di compleanno ai clienti.'
)

doc.add_paragraph()
doc.add_heading('9.1 Template WhatsApp', level=2)
doc.add_paragraph(
    'Crea template riutilizzabili per i messaggi WhatsApp:'
)

add_table_2col(
    ['Tipo Template', 'Uso'],
    [
        ('Promemoria', 'Reminder per appuntamenti'),
        ('Compleanno', 'Auguri di buon compleanno'),
        ('Promozione', 'Offerte e promozioni'),
        ('Libero', 'Messaggio personalizzato'),
    ]
)

doc.add_paragraph()
doc.add_heading('9.2 Template Email', level=2)
doc.add_paragraph(
    'Funziona come i template WhatsApp, ma con l\'aggiunta del campo Oggetto (Subject).'
)

doc.add_paragraph()
doc.add_heading('9.3 Creare un Template', level=2)
add_numbered('Clicca "+ Crea Template"')
add_numbered('Passo 1: inserisci nome, tipo (promemoria/compleanno/promozione/libero) e canale (WhatsApp/Email)')
add_numbered('Passo 2: scrivi il messaggio usando i segnaposto disponibili')
add_numbered('Clicca "Salva"')

doc.add_paragraph()
doc.add_heading('9.4 Segnaposto Disponibili', level=2)
doc.add_paragraph(
    'I segnaposto vengono sostituiti automaticamente con i dati reali quando invii il messaggio:'
)
add_table_2col(
    ['Segnaposto', 'Viene sostituito con'],
    [
        ('{nome}', 'Nome del cliente'),
        ('{cognome}', 'Cognome del cliente'),
        ('{data_appuntamento}', 'Data dell\'appuntamento'),
        ('{ora_appuntamento}', 'Ora dell\'appuntamento'),
        ('{trattamento}', 'Nome del trattamento'),
        ('{nome_centro}', 'Nome del centro estetico'),
    ]
)
add_tip('Clicca sui bottoni segnaposto per inserirli direttamente nel punto del cursore.')

doc.add_paragraph()
doc.add_heading('9.5 Compleanni', level=2)
doc.add_paragraph(
    'La sezione mostra i compleanni di oggi e quelli dei prossimi 7 giorni, '
    'con la possibilità di inviare auguri direttamente via WhatsApp o email.'
)

section_break()

# ═══════════════════════════════════════════════════════════════
# 10. REPORT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Report — Analisi Dati', level=1)
doc.add_paragraph(
    'La sezione Report offre analisi dettagliate sull\'andamento del centro.'
)

doc.add_paragraph()
doc.add_heading('10.1 Tab Panoramica', level=2)

doc.add_heading('Riepilogo Giornata', level=3)
doc.add_paragraph(
    'Card in evidenza con i dati del giorno corrente: incasso, numero appuntamenti e clienti unici.'
)

doc.add_heading('Confronto Mensile', level=3)
doc.add_paragraph(
    'Confronta il mese corrente con il precedente su tre metriche:'
)
add_bullet('Fatturato: importo e variazione percentuale')
add_bullet('Appuntamenti: conteggio e variazione')
add_bullet('Clienti: conteggio e variazione')

doc.add_heading('Panoramica Annuale', level=3)
doc.add_paragraph(
    'Riepilogo dell\'anno in corso: fatturato totale, appuntamenti totali, nuovi clienti '
    'e tasso di completamento.'
)

doc.add_heading('Top 5 Trattamenti e Clienti', level=3)
doc.add_paragraph(
    'Classifiche dei servizi più redditizi e dei clienti migliori, con barre grafiche proporzionali.'
)

doc.add_heading('Trend Mensile', level=3)
doc.add_paragraph(
    'Grafico a barre del fatturato degli ultimi 6 mesi per visualizzare l\'andamento.'
)

doc.add_paragraph()
doc.add_heading('10.2 Tab Interrogazione Libera', level=2)
doc.add_paragraph(
    'Strumento avanzato per analisi personalizzate:'
)
add_numbered('Seleziona il periodo (da — a)')
add_numbered('Filtra per cliente, trattamento, operatrice o stato')
add_numbered('Visualizza i risultati nella tabella')
add_numbered('Esporta in CSV per elaborazioni esterne (es. Excel)')
add_tip('L\'interrogazione libera è utile per estrarre dati specifici, '
        'ad esempio "tutti gli appuntamenti di un cliente nell\'ultimo anno".')

section_break()

# ═══════════════════════════════════════════════════════════════
# 11. INSIGHTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Insights — Business Intelligence', level=1)
doc.add_paragraph(
    'La sezione Insights analizza automaticamente i dati del gestionale e genera '
    'suggerimenti actionable per migliorare il business. Non sono solo numeri: '
    'sono consigli pratici con priorità e azioni consigliate.'
)

doc.add_paragraph()
doc.add_heading('11.1 Tab Insights — Messaggi Intelligenti', level=2)
doc.add_paragraph(
    'Card colorate ordinate per priorità con consigli generati automaticamente dall\'analisi dei dati:'
)
add_table_3col(
    ['Priorità', 'Colore', 'Significato'],
    [
        ('Alta', 'Rosso', 'Richiede attenzione immediata'),
        ('Media', 'Arancione', 'Da valutare a breve'),
        ('Bassa', 'Verde', 'Informativo, buone notizie o suggerimenti'),
    ]
)
doc.add_paragraph('Esempi di insight generati:')
add_bullet('Clienti a rischio: "X clienti non tornano da oltre 60 giorni — contattali"')
add_bullet('Scorte in esaurimento: "X prodotti finiranno entro 2 settimane"')
add_bullet('Trend fatturato: "Fatturato in calo del X% rispetto al mese precedente"')
add_bullet('Giorno debole: "Il martedì ha il 40% in meno di appuntamenti — valuta una promozione"')
add_bullet('Trattamento top: "Il trattamento X è il più redditizio con margine del X%"')
add_bullet('Tasso di ritorno: "Solo il X% dei clienti torna dopo la prima visita"')

add_tip('Ogni card ha una tendina espandibile che spiega il dato nel dettaglio e suggerisce l\'azione da intraprendere.')

doc.add_paragraph()
doc.add_heading('11.2 Tab Clienti — Segmentazione', level=2)
doc.add_paragraph(
    'Suddivide automaticamente i clienti in 5 segmenti basati sul comportamento reale:'
)

add_table_3col(
    ['Segmento', 'Criterio', 'Azione Consigliata'],
    [
        ('VIP', 'Spesa > 2x la media + almeno 5 visite', 'Fidelizza con trattamento esclusivo'),
        ('Abituali', 'Attivi negli ultimi 60 giorni', 'Mantieni la relazione'),
        ('Nuovi', '1 sola visita negli ultimi 60 giorni', 'Follow-up per fidelizzare'),
        ('A Rischio', 'Assenti da 60-120 giorni', 'Contatta prima che li perdi'),
        ('Persi', 'Assenti da oltre 120 giorni', 'Campagna di riattivazione'),
    ]
)
doc.add_paragraph()
doc.add_paragraph('KPI in evidenza:')
add_bullet('Clienti Attivi: numero totale di clienti con almeno un appuntamento')
add_bullet('Tasso di Ritorno: percentuale di clienti che tornano dopo la prima visita')
add_bullet('Clienti VIP: numero di clienti nel segmento VIP')
doc.add_paragraph()
doc.add_paragraph(
    'Ogni segmento è espandibile e mostra la lista dei clienti con: nome, numero visite, '
    'spesa totale e data ultima visita.'
)

doc.add_paragraph()
doc.add_heading('11.3 Tab Performance', level=2)

doc.add_heading('Heatmap Occupazione', level=3)
doc.add_paragraph(
    'Griglia visuale che mostra i giorni e le fasce orarie più occupate degli ultimi 90 giorni. '
    'Le celle più scure indicano maggiore affluenza.'
)
add_bullet('Righe: fasce orarie (09-10, 10-11, … 18-19)')
add_bullet('Colonne: giorni della settimana (Lun — Sab)')
add_tip('Identifica le fasce vuote per proporre promozioni mirate.')

doc.add_heading('Margini per Trattamento', level=3)
doc.add_paragraph(
    'Classifica dei trattamenti per redditività reale (ultimi 12 mesi):'
)
add_bullet('Ricavo totale del trattamento')
add_bullet('Costo prodotti consumati (basato sui movimenti magazzino)')
add_bullet('Margine = Ricavo - Costo')
add_bullet('Margine percentuale')
add_note('Per avere margini accurati, registra sempre i prodotti usati negli appuntamenti completati.')

doc.add_heading('Confronto Mensile', level=3)
doc.add_paragraph(
    'Grafico a barre degli ultimi 12 mesi con: fatturato, numero appuntamenti, clienti unici e ticket medio.'
)

doc.add_paragraph()
doc.add_heading('11.4 Tab Previsioni', level=2)

doc.add_heading('Previsione Fatturato', level=3)
doc.add_paragraph(
    'Stima del fatturato del prossimo mese basata sulla media degli ultimi 3 mesi, '
    'con dettaglio dei singoli mesi usati nel calcolo.'
)

doc.add_heading('Scorte in Esaurimento', level=3)
doc.add_paragraph(
    'Lista dei prodotti che finiranno entro 30 giorni al ritmo attuale di consumo:'
)
add_bullet('Nome prodotto e giacenza attuale')
add_bullet('Consumo medio giornaliero (calcolato sugli ultimi 90 giorni)')
add_bullet('Giorni rimanenti stimati')
add_bullet('Codice colore: rosso (< 7 giorni), arancione (7-14 giorni), giallo (14-30 giorni)')

section_break()

# ═══════════════════════════════════════════════════════════════
# 12. IMPOSTAZIONI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Impostazioni — Configurazione', level=1)

doc.add_heading('12.1 Il Mio Account', level=2)
doc.add_paragraph('Gestione del profilo utente:')
add_bullet('Cambio password: inserisci la password corrente, la nuova e conferma')
add_bullet('Parola chiave di recupero: usata per il ripristino dell\'account')

doc.add_paragraph()
doc.add_heading('12.2 Aspetto', level=2)
doc.add_paragraph('Personalizzazione dell\'interfaccia:')
add_bullet('Palette colori: scegli tra diverse combinazioni di colori')
add_bullet('Tema: Chiaro, Scuro o Automatico (segue le impostazioni del sistema operativo)')
add_tip('Il cambio tema è immediato, non serve riavviare.')

doc.add_paragraph()
doc.add_heading('12.3 Dati Azienda', level=2)
doc.add_paragraph(
    'Dati del centro estetico usati in tutto il gestionale:'
)
add_table_2col(
    ['Campo', 'Dove viene usato'],
    [
        ('Nome centro', 'Template messaggi {nome_centro}, intestazioni'),
        ('Indirizzo / Città / CAP', 'Informazioni di contatto'),
        ('Telefono / Email', 'Informazioni di contatto'),
        ('P.IVA', 'Dati fiscali'),
        ('Orario apertura/chiusura', 'Calcolo slot liberi, griglia agenda'),
        ('Durata slot (minuti)', 'Griglia oraria dell\'agenda'),
        ('Giorni lavorativi', 'Calcolo saturazione settimanale'),
    ]
)
add_note('Gli orari di apertura/chiusura sono fondamentali per il calcolo corretto degli slot liberi in Dashboard.')

doc.add_paragraph()
doc.add_heading('12.4 Email SMTP', level=2)
doc.add_paragraph(
    'Configura il server SMTP per l\'invio di email (promemoria, auguri, ecc.):'
)
add_numbered('Inserisci Host, Porta, Username e Password del server SMTP')
add_numbered('Imposta l\'email e il nome mittente')
add_numbered('Seleziona il tipo di crittografia (TLS o SSL)')
add_numbered('Abilita il servizio con il toggle')
add_numbered('Clicca "Test Connessione" per verificare che funzioni')
add_numbered('Salva la configurazione')

add_tip('Per Gmail: host=smtp.gmail.com, porta=587, crittografia=TLS. '
        'Serve una "Password per le app" (non la password del tuo account).')

doc.add_paragraph()
doc.add_heading('12.5 Licenza', level=2)
doc.add_paragraph(
    'Visualizza lo stato della licenza e il codice di attivazione.'
)

doc.add_paragraph()
doc.add_heading('12.6 Gestione Utenti', level=2)
doc.add_paragraph(
    'Crea e gestisci gli utenti che possono accedere al gestionale, con ruoli e permessi diversi.'
)

doc.add_paragraph()
doc.add_heading('12.7 Backup e Ripristino', level=2)
doc.add_paragraph(
    'Fondamentale per proteggere i dati del centro:'
)

doc.add_heading('Creare un Backup', level=3)
add_numbered('Clicca "Crea Backup"')
add_numbered('Aggiungi una descrizione (es. "Backup settimanale")')
add_numbered('Conferma: il file viene salvato nella cartella backup dell\'app')

doc.add_heading('Ripristinare un Backup', level=3)
add_numbered('Nella lista backup, individua quello desiderato')
add_numbered('Clicca "Ripristina"')
add_numbered('Conferma nel dialog di avviso (il ripristino sovrascrive i dati attuali)')
add_note('Il ripristino è un\'operazione irreversibile. Crea un backup dei dati attuali prima di ripristinare.')

doc.add_heading('Altre azioni', level=3)
add_bullet('Scarica backup: salva una copia in una posizione a tua scelta')
add_bullet('Elimina backup: rimuove il file di backup')
add_bullet('Apri cartella backup: apre la cartella nel file manager')

add_tip('Consiglio: crea un backup almeno una volta a settimana e conserva una copia '
        'su un disco esterno o in cloud.')

doc.add_paragraph()
doc.add_heading('12.8 Aggiornamenti', level=2)
doc.add_paragraph(
    'Gestione degli aggiornamenti automatici (OTA — Over The Air):'
)
add_numbered('Clicca "Controlla aggiornamenti"')
add_numbered('Se disponibile, clicca "Scarica aggiornamento" — una barra di progresso mostra lo stato')
add_numbered('Al termine, clicca "Riavvia per aggiornare" — l\'app si riavvia con la nuova versione')
add_tip('Gli aggiornamenti sono incrementali e non cancellano i tuoi dati.')

section_break()

# ═══════════════════════════════════════════════════════════════
# 13. AGGIORNAMENTI OTA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Aggiornamenti — OTA Updates', level=1)
doc.add_paragraph(
    'Beauty Manager Pro supporta gli aggiornamenti automatici over-the-air (OTA). '
    'Quando una nuova versione è disponibile, l\'app ti avvisa e puoi aggiornare con un click '
    'dalla sezione Impostazioni → Aggiornamenti.'
)
doc.add_paragraph(
    'Il processo è sicuro: i tuoi dati non vengono mai toccati durante l\'aggiornamento. '
    'Solo i file dell\'applicazione vengono sostituiti.'
)

section_break()

# ═══════════════════════════════════════════════════════════════
# 14. SCORCIATOIE E TRUCCHI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. Scorciatoie e Trucchi', level=1)

doc.add_heading('14.1 Navigazione Rapida', level=2)
add_bullet('Clicca sulla card "Appuntamenti Oggi" in Dashboard per aprire l\'Agenda sulla data odierna')
add_bullet('Clicca su un cliente ovunque nel programma per aprire la sua scheda')
add_bullet('Usa la barra di ricerca in ogni sezione per trovare velocemente ciò che cerchi')

doc.add_paragraph()
doc.add_heading('14.2 Agenda', level=2)
add_bullet('Trascina un appuntamento per spostarlo a un nuovo orario')
add_bullet('Trascina tra colonne per cambiare operatrice')
add_bullet('Ridimensiona dal bordo inferiore per cambiare durata')
add_bullet('Clicca su uno slot vuoto per creare un appuntamento precompilato')

doc.add_paragraph()
doc.add_heading('14.3 Prodotti', level=2)
add_bullet('Nella sezione Prodotti Usati (appuntamento completato), puoi scansionare il codice a barre '
           'premendo Enter nel campo di ricerca')
add_bullet('I prodotti vengono scaricati automaticamente dal magazzino al salvataggio')

doc.add_paragraph()
doc.add_heading('14.4 Template Messaggi', level=2)
add_bullet('Usa i pulsanti segnaposto per inserirli nel messaggio senza doverli digitare')
add_bullet('L\'anteprima mostra il messaggio con dati di esempio, aggiornata in tempo reale')

doc.add_paragraph()
doc.add_heading('14.5 Buone Pratiche', level=2)
add_numbered('Fai un backup settimanale (Impostazioni → Backup)')
add_numbered('Completa gli appuntamenti a fine giornata per avere statistiche accurate')
add_numbered('Registra i prodotti usati negli appuntamenti per i calcoli dei margini')
add_numbered('Controlla la sezione Insights periodicamente per identificare opportunità')
add_numbered('Mantieni aggiornati i dati dei clienti (consensi, contatti) per le comunicazioni')

# ═══════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Beauty Manager Pro — Manuale Utente v0.1.0')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ────────────────────────────────────────────────────
output_path = os.path.expanduser('~/Apps/Beauty Manager/Beauty_Manager_Pro_Manuale_Utente.docx')
doc.save(output_path)
print(f"Manuale salvato in: {output_path}")
